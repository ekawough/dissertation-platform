import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_chapter_docx(title: str, content: str, client_name: str, institution: str) -> bytes:
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    # Default style — Times New Roman 12pt double-spaced left-aligned
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    from docx.shared import Pt as P
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ragged right

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_para.add_run(client_name)

    inst_para = doc.add_paragraph()
    inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_para.add_run(institution)

    doc.add_page_break()

    # Content — parse markdown-like headings
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].upper())
            run.bold = True
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[3:])
            run.bold = True
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[4:])
            run.bold = True
            run.italic = True
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Double spacing
            from docx.shared import Pt as Pt2
            p.paragraph_format.space_after = Pt2(0)
            from docx.oxml.ns import qn as qn2
            pPr = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn2('w:line'), '480')
            spacing.set(qn2('w:lineRule'), 'auto')
            pPr.append(spacing)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def export_full_dissertation_docx(client_name: str, institution: str, topic: str, chapters: list) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.add_run(topic.upper()).bold = True
    doc.add_paragraph()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(client_name)
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run(institution)
    doc.add_page_break()

    for ch in chapters:
        lines = (ch.get("content") or "").split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(line[2:].upper()).bold = True
            elif line.startswith('## '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(line[3:]).bold = True
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
